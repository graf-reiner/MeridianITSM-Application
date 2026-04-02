"""Generate OAuth Email Setup Guide as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    hs.font.name = 'Calibri'
    if level == 1:
        hs.font.size = Pt(24)
        hs.paragraph_format.space_before = Pt(24)
    elif level == 2:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(18)
    else:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(12)


def add_screenshot_box(doc, caption):
    """Add a styled placeholder box for a screenshot."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    # Style the cell
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F4F8',
    })
    shading.append(shading_elm)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n📷  SCREENSHOT: {caption}\n\n')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
    run.font.italic = True
    run2 = p.add_run('(Insert screenshot here)\n')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    run2.font.italic = True
    doc.add_paragraph()  # spacing


def add_code_block(doc, text):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
    # Add background shading to paragraph
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F1F5F9',
    })
    pPr.append(shd)


def add_note(doc, text):
    """Add an info/note callout."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('ℹ️  ')
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)


def add_warning(doc, text):
    """Add a warning callout."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('⚠️  ')
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xb4, 0x53, 0x09)


# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════

doc.add_paragraph('\n\n\n')
title = doc.add_heading('MeridianITSM', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(36)
title.runs[0].font.color.rgb = RGBColor(0x02, 0x84, 0xc7)

subtitle = doc.add_heading('OAuth Email Setup Guide', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(22)
subtitle.runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)
subtitle.paragraph_format.space_before = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Google Workspace & Gmail  •  Microsoft 365 & Outlook')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

doc.add_paragraph('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0  •  April 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1', 'Prerequisites'),
    ('2', 'Part 1: Google (Gmail & Google Workspace)'),
    ('2.1', '  Step 1: Create a Google Cloud Project'),
    ('2.2', '  Step 2: Enable the Gmail API'),
    ('2.3', '  Step 3: Configure the OAuth Consent Screen'),
    ('2.4', '  Step 4: Create OAuth 2.0 Credentials'),
    ('2.5', '  Step 5: Add Credentials to MeridianITSM'),
    ('2.6', '  Step 6: Connect a Google Email Account'),
    ('3', 'Part 2: Microsoft (Outlook & Microsoft 365)'),
    ('3.1', '  Step 1: Register an Application in Entra ID'),
    ('3.2', '  Step 2: Create a Client Secret'),
    ('3.3', '  Step 3: Configure API Permissions'),
    ('3.4', '  Step 4: Add Credentials to MeridianITSM'),
    ('3.5', '  Step 5: Connect a Microsoft Email Account'),
    ('4', 'Troubleshooting'),
    ('5', 'Environment Variable Summary'),
]
for num, label in toc_items:
    p = doc.add_paragraph()
    if '.' not in num:
        run = p.add_run(f'{label}')
        run.font.bold = True
        run.font.size = Pt(11)
    else:
        run = p.add_run(f'{label}')
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PREREQUISITES
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Prerequisites', level=1)
doc.add_paragraph('Before setting up OAuth email accounts, ensure you have:')
items = [
    'MeridianITSM running with the API accessible at a public URL (e.g., https://meridian.cybordyne.net)',
    'Admin access to Google Cloud Console and/or Microsoft Entra ID (Azure AD)',
    'Admin role in MeridianITSM',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
add_note(doc, 'The OAuth callback URL must be publicly accessible. OAuth providers redirect the user\'s browser to this URL after authorization.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 1: GOOGLE
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Part 1: Google (Gmail & Google Workspace)', level=1)
doc.add_paragraph('This section guides you through creating a Google Cloud OAuth application and connecting a Gmail or Google Workspace email account to MeridianITSM.')

# ── Step 1 ──
doc.add_heading('Step 1: Create a Google Cloud Project', level=2)
doc.add_paragraph('1.  Go to Google Cloud Console: https://console.cloud.google.com/')
doc.add_paragraph('2.  Click the project dropdown in the top-left corner, then click New Project.')
add_screenshot_box(doc, 'Google Cloud Console — Project dropdown → New Project button')
doc.add_paragraph('3.  Enter a project name (e.g., "MeridianITSM-Email") and click Create.')
add_screenshot_box(doc, 'New Project dialog — Name field and Create button')
doc.add_paragraph('4.  Select the newly created project from the project dropdown.')

# ── Step 2 ──
doc.add_heading('Step 2: Enable the Gmail API', level=2)
doc.add_paragraph('1.  In the left sidebar, navigate to APIs & Services → Library.')
doc.add_paragraph('2.  Search for "Gmail API" in the search bar.')
doc.add_paragraph('3.  Click on Gmail API in the results, then click the Enable button.')
add_screenshot_box(doc, 'Gmail API page — Enable button highlighted')
add_note(doc, 'The Gmail API must be enabled for OAuth to request mail access scopes.')

# ── Step 3 ──
doc.add_heading('Step 3: Configure the OAuth Consent Screen', level=2)
doc.add_paragraph('1.  Navigate to APIs & Services → OAuth consent screen.')
doc.add_paragraph('2.  Choose a user type:')

tbl = doc.add_table(rows=3, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0]
for i, text in enumerate(['User Type', 'When to Use']):
    hdr.cells[i].text = text
    for p in hdr.cells[i].paragraphs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
tbl.rows[1].cells[0].text = 'Internal'
tbl.rows[1].cells[1].text = 'All email accounts are within your Google Workspace organization. No Google review needed.'
tbl.rows[2].cells[0].text = 'External'
tbl.rows[2].cells[1].text = 'You need to connect personal Gmail accounts. Requires Google review for production use.'
doc.add_paragraph()

add_screenshot_box(doc, 'OAuth consent screen — User type selection (Internal vs External)')

doc.add_paragraph('3.  Click Create and fill in the required fields:')
for item in [
    'App name: MeridianITSM',
    'User support email: your admin email address',
    'Developer contact email: your admin email address',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('4.  Click Save and Continue.')
doc.add_paragraph('5.  On the Scopes page, click Add or Remove Scopes and add the following:')

tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0]
for i, text in enumerate(['Scope', 'Purpose']):
    hdr.cells[i].text = text
    for p in hdr.cells[i].paragraphs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
scopes = [
    ('https://mail.google.com/', 'Full Gmail access (IMAP/SMTP)'),
    ('openid', 'OpenID Connect identity'),
    ('email', 'Read user email address'),
    ('profile', 'Read user name'),
]
for i, (scope, purpose) in enumerate(scopes):
    tbl.rows[i+1].cells[0].text = scope
    tbl.rows[i+1].cells[1].text = purpose
doc.add_paragraph()

add_screenshot_box(doc, 'Add Scopes dialog — mail.google.com scope selected')

doc.add_paragraph('6.  Click Update, then Save and Continue.')
doc.add_paragraph('7.  On the Test Users page (External type only): add the Gmail accounts you\'ll use during testing.')
doc.add_paragraph('8.  Click Save and Continue, then Back to Dashboard.')

# ── Step 4 ──
doc.add_heading('Step 4: Create OAuth 2.0 Credentials', level=2)
doc.add_paragraph('1.  Navigate to APIs & Services → Credentials.')
doc.add_paragraph('2.  Click + Create Credentials → OAuth client ID.')
add_screenshot_box(doc, 'Credentials page — Create Credentials dropdown → OAuth client ID')

doc.add_paragraph('3.  Configure the OAuth client:')
for item in [
    'Application type: Web application',
    'Name: MeridianITSM Email',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('4.  Under Authorized redirect URIs, click + Add URI and enter:')
add_code_block(doc, 'https://meridian.cybordyne.net/api/v1/email-accounts/oauth/callback')

add_warning(doc, 'Replace with your actual domain. This must exactly match your APP_URL environment variable.')
add_screenshot_box(doc, 'OAuth client ID config — Redirect URI field filled in')

doc.add_paragraph('5.  Click Create.')
doc.add_paragraph('6.  A dialog will show your Client ID and Client Secret. Copy both values — you\'ll need them in the next step.')
add_screenshot_box(doc, 'OAuth client created dialog — Client ID and Client Secret displayed')

add_warning(doc, 'Save the Client Secret now — you cannot view it again later. You\'d have to create a new one.')

# ── Step 5 ──
doc.add_heading('Step 5: Add Credentials to MeridianITSM', level=2)
doc.add_paragraph('SSH into your server and edit the API environment file:')
add_code_block(doc, 'ssh meridian-dev\nnano /opt/meridian/apps/api/.env')
doc.add_paragraph('Add these lines:')
add_code_block(doc, 'GOOGLE_CLIENT_ID=123456789-xxxxxxx.apps.googleusercontent.com\nGOOGLE_CLIENT_SECRET=GOCSPX-xxxxxxxxxxxxxxxx\nAPP_URL=https://meridian.cybordyne.net')
doc.add_paragraph('Save the file and restart the API:')
add_code_block(doc, 'cd /opt/meridian && pm2 restart api')

# ── Step 6 ──
doc.add_heading('Step 6: Connect a Google Email Account in MeridianITSM', level=2)
steps = [
    'Log in to MeridianITSM as an admin.',
    'Navigate to Settings → Email Accounts.',
    'Click Add Account.',
    'Select Google (Workspace & Gmail) from the provider selection.',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'{i+1}.  {step}')

add_screenshot_box(doc, 'MeridianITSM — Provider selection dialog showing Google, Microsoft, and Manual options')

steps2 = [
    'A popup window opens with Google\'s sign-in page. Sign in with the Google account you want to use for email.',
    'Grant the requested permissions when prompted.',
]
for i, step in enumerate(steps2, start=5):
    doc.add_paragraph(f'{i}.  {step}')

add_screenshot_box(doc, 'Google OAuth consent — Permission grant dialog')

steps3 = [
    'The popup closes automatically and returns you to MeridianITSM.',
    'Configure the account settings:',
]
for i, step in enumerate(steps3, start=7):
    doc.add_paragraph(f'{i}.  {step}')

for item in [
    'Display Name — e.g., "Support Inbox"',
    'Poll Interval — how often to check for new emails (default: 5 minutes)',
    'Default Queue — which ticket queue new emails create tickets in',
    'Default Category — optional category for new tickets',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('9.  Click Save. The account is now active and polling for emails.')
add_screenshot_box(doc, 'MeridianITSM — Post-connect configuration form')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# PART 2: MICROSOFT
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Part 2: Microsoft (Outlook & Microsoft 365)', level=1)
doc.add_paragraph('This section guides you through registering an application in Microsoft Entra ID (Azure AD) and connecting an Outlook or Microsoft 365 mailbox.')

# ── Step 1 ──
doc.add_heading('Step 1: Register an Application in Microsoft Entra ID', level=2)
doc.add_paragraph('1.  Go to the Microsoft Entra Admin Center: https://entra.microsoft.com/')
doc.add_paragraph('2.  In the left sidebar, navigate to Identity → Applications → App registrations.')
doc.add_paragraph('3.  Click + New registration.')
add_screenshot_box(doc, 'Entra ID — App registrations page → New registration button')

doc.add_paragraph('4.  Fill in the registration form:')
for item in [
    'Name: MeridianITSM Email',
    'Supported account types: Accounts in any organizational directory (Multitenant)',
]:
    doc.add_paragraph(item, style='List Bullet')

add_note(doc, 'Choose "Single tenant" if all mailboxes are within your one organization.')

doc.add_paragraph('5.  Under Redirect URI:')
for item in [
    'Platform: Web',
    'URI: https://meridian.cybordyne.net/api/v1/email-accounts/oauth/callback',
]:
    doc.add_paragraph(item, style='List Bullet')

add_screenshot_box(doc, 'Register an application form — Name, account type, and redirect URI fields')

doc.add_paragraph('6.  Click Register.')
doc.add_paragraph('7.  On the Overview page, copy the Application (client) ID. You\'ll need this later.')
add_screenshot_box(doc, 'App registration Overview — Application (client) ID highlighted')

# ── Step 2 ──
doc.add_heading('Step 2: Create a Client Secret', level=2)
doc.add_paragraph('1.  In your app registration, go to Certificates & secrets in the left sidebar.')
doc.add_paragraph('2.  Click + New client secret.')
doc.add_paragraph('3.  Enter a description (e.g., "MeridianITSM") and choose an expiration period (recommended: 24 months).')
doc.add_paragraph('4.  Click Add.')
add_screenshot_box(doc, 'Certificates & secrets — New client secret dialog')

doc.add_paragraph('5.  Copy the secret Value immediately.')
add_warning(doc, 'The secret value is only shown once. If you navigate away, you cannot retrieve it — you\'d need to create a new one.')
add_screenshot_box(doc, 'Client secret created — Value column highlighted (copy this)')

# ── Step 3 ──
doc.add_heading('Step 3: Configure API Permissions', level=2)
doc.add_paragraph('1.  In the left sidebar, go to API permissions.')
doc.add_paragraph('2.  Click + Add a permission.')
doc.add_paragraph('3.  Select the APIs my organization uses tab.')
doc.add_paragraph('4.  Search for "Office 365 Exchange Online" and select it.')
add_screenshot_box(doc, 'Request API permissions — Searching for Office 365 Exchange Online')

doc.add_paragraph('5.  Choose Delegated permissions and add:')

tbl = doc.add_table(rows=3, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0]
for i, text in enumerate(['Permission', 'Purpose']):
    hdr.cells[i].text = text
    for p in hdr.cells[i].paragraphs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)
tbl.rows[1].cells[0].text = 'IMAP.AccessAsUser.All'
tbl.rows[1].cells[1].text = 'Read emails via IMAP'
tbl.rows[2].cells[0].text = 'SMTP.Send'
tbl.rows[2].cells[1].text = 'Send emails via SMTP'
doc.add_paragraph()

add_screenshot_box(doc, 'Delegated permissions — IMAP.AccessAsUser.All and SMTP.Send selected')

doc.add_paragraph('6.  Click Add permissions.')
doc.add_paragraph('7.  Ensure these Microsoft Graph permissions are also present (usually added by default):')
for item in ['openid', 'email', 'profile', 'offline_access (required for refresh tokens)']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('8.  If you are a tenant admin, click Grant admin consent for [your organization].')
add_screenshot_box(doc, 'API permissions list — Grant admin consent button highlighted')

add_note(doc, 'If you cannot grant admin consent, each user will be prompted to consent individually during the OAuth flow.')

# ── Step 4 ──
doc.add_heading('Step 4: Add Credentials to MeridianITSM', level=2)
doc.add_paragraph('SSH into your server and edit the API environment file:')
add_code_block(doc, 'ssh meridian-dev\nnano /opt/meridian/apps/api/.env')
doc.add_paragraph('Add these lines:')
add_code_block(doc, 'MICROSOFT_CLIENT_ID=xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx\nMICROSOFT_CLIENT_SECRET=xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx\nAPP_URL=https://meridian.cybordyne.net')
add_note(doc, 'If you already set APP_URL for Google, don\'t add it twice.')
doc.add_paragraph('Save and restart the API:')
add_code_block(doc, 'cd /opt/meridian && pm2 restart api')

# ── Step 5 ──
doc.add_heading('Step 5: Connect a Microsoft Email Account in MeridianITSM', level=2)
steps = [
    'Log in to MeridianITSM as an admin.',
    'Navigate to Settings → Email Accounts.',
    'Click Add Account.',
    'Select Microsoft 365 (Outlook & Exchange).',
    'A popup window opens with Microsoft\'s sign-in page. Sign in with the Microsoft account/mailbox you want to use.',
    'Grant the requested permissions when prompted.',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'{i+1}.  {step}')

add_screenshot_box(doc, 'Microsoft OAuth consent — Permissions requested dialog')

steps2 = [
    'The popup closes automatically and returns you to MeridianITSM.',
    'Configure the account: Display Name, Poll Interval, Default Queue, Default Category.',
    'Click Save. The account is now active and polling for emails.',
]
for i, step in enumerate(steps2, start=7):
    doc.add_paragraph(f'{i}.  {step}')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Troubleshooting', level=1)

issues = [
    ('OAuth popup doesn\'t open or gets blocked',
     'Ensure your browser allows popups from your MeridianITSM domain. Check that APP_URL in the .env file matches the exact domain you\'re accessing (including https://).'),
    ('"redirect_uri_mismatch" error (Google)',
     'The redirect URI in Google Cloud Console must exactly match:\nhttps://your-domain.com/api/v1/email-accounts/oauth/callback\nNo trailing slash. Must use https://.'),
    ('"AADSTS50011: The redirect URI does not match" (Microsoft)',
     'Same as above — the redirect URI in Entra ID must exactly match the callback URL. Ensure the platform is set to "Web" (not SPA).'),
    ('Account shows "Disconnected" after working',
     'The refresh token was revoked (user changed password, admin revoked access, or token expired). Click Reconnect on the email account to re-authorize. MeridianITSM automatically deactivates accounts when refresh fails and notifies the admin.'),
    ('Google says "This app isn\'t verified"',
     'Expected during development with External user type. Click Advanced → Go to MeridianITSM (unsafe) to proceed. For production, submit your app for Google verification via the consent screen page.'),
    ('Microsoft says "Need admin approval"',
     'Your Entra ID tenant requires admin consent for the requested permissions. A tenant admin must click Grant admin consent in the app\'s API permissions page.'),
    ('Emails not being received',
     'Check that the account is Active in Settings → Email Accounts. Check worker logs:\nssh meridian-dev "pm2 logs worker --lines 50"\nLook for [email-inbound] log entries.'),
    ('Emails not being sent',
     'Check API logs: ssh meridian-dev "pm2 logs api --lines 50"\nFor Google: ensure the Gmail API is enabled. For Microsoft: ensure SMTP.Send permission was granted with admin consent.'),
]

for title, body in issues:
    doc.add_heading(title, level=3)
    doc.add_paragraph(body)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# ENVIRONMENT VARIABLE SUMMARY
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Environment Variable Summary', level=1)
doc.add_paragraph('Add all of these to /opt/meridian/apps/api/.env on the server:')
add_code_block(doc, '''# Required for OAuth to work
APP_URL=https://meridian.cybordyne.net

# Google OAuth (only if using Google email)
GOOGLE_CLIENT_ID=your-client-id.apps.googleusercontent.com
GOOGLE_CLIENT_SECRET=your-client-secret

# Microsoft OAuth (only if using Microsoft email)
MICROSOFT_CLIENT_ID=your-application-id
MICROSOFT_CLIENT_SECRET=your-client-secret

# Already present — used for encrypting OAuth tokens
ENCRYPTION_KEY=0123456789abcdef...''')

doc.add_paragraph()
doc.add_paragraph('After adding or changing any values, restart the API and worker:')
add_code_block(doc, 'ssh meridian-dev "cd /opt/meridian && pm2 restart api worker"')

# ── Save ────────────────────────────────────────────────────────────────────
output_path = 'docs/MeridianITSM-OAuth-Email-Setup-Guide.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
